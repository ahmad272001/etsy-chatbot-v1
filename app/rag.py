import os
import uuid
from typing import List, Optional, Tuple
import PyPDF2
from qdrant_client import QdrantClient
from qdrant_client.http.models import PointStruct, Filter, FieldCondition, MatchValue
from app.config import settings
from app.models import RetrievalRef
from openai import OpenAI


class RAGEngine:
    def __init__(self):
        """Initialize RAG engine with Qdrant and OpenAI."""
        # Initialize Qdrant client
        self.qdrant_client = QdrantClient(
            url=settings.qdrant_url,  # e.g. "https://<cluster-id>.<region>.gcp.cloud.qdrant.io"
            api_key=settings.qdrant_api_key
        )

        # Ensure collection exists with proper indexing
        try:
            collections = self.qdrant_client.get_collections().collections
            if not any(c.name == settings.qdrant_collection_name for c in collections):
                # Create collection with proper indexing
                self.qdrant_client.recreate_collection(
                    collection_name=settings.qdrant_collection_name,
                    vectors_config={
                        "size": 3072,  # OpenAI text-embedding-3-large dimensions
                        "distance": "Cosine"
                    },
                    optimizers_config={
                        "default_segment_number": 2,
                        "memmap_threshold": 10000
                    }
                )
                
                # Create index on source_file field for filtering
                self.qdrant_client.create_payload_index(
                    collection_name=settings.qdrant_collection_name,
                    field_name="source_file",
                    field_schema="keyword"
                )
                
                pass
            else:
                # Ensure index exists on source_file field
                try:
                    self.qdrant_client.create_payload_index(
                        collection_name=settings.qdrant_collection_name,
                        field_name="source_file",
                        field_schema="keyword"
                    )
                except Exception as index_error:
                    # Index might already exist
                    pass
                    
        except Exception as e:
            print(f"❌ Error initializing Qdrant: {e}")
            raise

        # Initialize OpenAI client
        self.openai_client = OpenAI(api_key=settings.openai_api_key)


    def load_document_from_upload(self, file_path: str, filename: str):
        """Load document content from a PDF or Word file."""
        if filename.lower().endswith(".pdf"):
            try:
                with open(file_path, "rb") as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text_content = ""
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text() or ""
                        text_content += page_text + "\n"
                    text_content = text_content.strip()
                    if text_content:
                        return {"id": filename, "text": text_content, "page_count": len(pdf_reader.pages)}
                    else:
                        return None
            except Exception as e:
                return None
        elif filename.lower().endswith(".docx"):
            try:
                from docx import Document
                doc = Document(file_path)
                text_content = ""
                for paragraph in doc.paragraphs:
                    text_content += paragraph.text + "\n"
                text_content = text_content.strip()
                if text_content:
                    return {"id": filename, "text": text_content, "page_count": len(doc.paragraphs) // 50 + 1}
                else:
                    return None
            except Exception as e:
                return None
        else:
            return None


    def split_text(self, text: str, chunk_size: int = 850, chunk_overlap: int = 100):
        """Split text into chunks with overlap."""
        chunks = []
        start = 0
        text = text.strip()
        while start < len(text):
            end = start + chunk_size
            if end < len(text):
                search_start = max(start, end - 100)
                last_period = text.rfind('.', search_start, end)
                last_exclamation = text.rfind('!', search_start, end)
                last_question = text.rfind('?', search_start, end)
                last_newline = text.rfind('\n', search_start, end)
                break_points = [last_period, last_exclamation, last_question, last_newline]
                valid_break_points = [p for p in break_points if p > start]
                if valid_break_points:
                    end = max(valid_break_points) + 1
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            start = end - chunk_overlap
        return chunks


    def preprocess_document(self, document, chunk_size: int = 1000, chunk_overlap: int = 20):
        """Split a single document into chunks."""
        chunks = self.split_text(document["text"], chunk_size, chunk_overlap)
        return [
            {"id": f"{document['id']}_chunk{i}", "text": chunk, "source_file": document["id"]}
            for i, chunk in enumerate(chunks, start=1)
        ]


    def get_openai_embedding(self, text: str):
        """Generate OpenAI embedding for text."""
        response = self.openai_client.embeddings.create(
            input=text,
            model=settings.embedding_model
        )
        return response.data[0].embedding


    def generate_embeddings(self, chunked_documents):
        """Generate embeddings for all chunks."""
        for doc in chunked_documents:
            doc["embedding"] = self.get_openai_embedding(doc["text"])
        return chunked_documents


    def add_documents_to_qdrant(self, chunked_documents):
        """Insert chunks with embeddings into Qdrant."""
        points = [
            PointStruct(
                id=str(uuid.uuid4()),
                vector=doc["embedding"],  # Use default vector field
                payload={
                    "text": doc["text"],
                    "source_file": doc["source_file"],
                    "chunk_id": doc["id"]
                }
            )
            for doc in chunked_documents
        ]
        self.qdrant_client.upsert(
            collection_name=settings.qdrant_collection_name,
            points=points
        )


    def add_document(self, file_path: str, filename: str) -> Tuple[str, int]:
        """Process and add a document to Qdrant."""
        document = self.load_document_from_upload(file_path, filename)
        if not document:
            raise ValueError("No content to add")
        chunks = self.preprocess_document(document)
        chunks = self.generate_embeddings(chunks)
        self.add_documents_to_qdrant(chunks)
        return str(uuid.uuid4()), document["page_count"]


    def query_documents(self, query: str, n_results: int = 2):
        """Search Qdrant for relevant chunks."""
        query_embedding = self.get_openai_embedding(query)
        search_result = self.qdrant_client.search(
            collection_name=settings.qdrant_collection_name,
            query_vector=query_embedding,
            limit=n_results
        )
        return [hit.payload["text"] for hit in search_result]


    def generate_response(self, question: str, relevant_chunks: List[str]):
        """Generate AI response with retrieved context."""
        context = "\n\n".join(relevant_chunks)
        prompt = (
            "You are an assistant for question-answering tasks. "
            "Use the retrieved context to answer the question concisely.\n\n"
            f"Context:\n{context}\n\nQuestion:\n{question}"
        )
        response = self.openai_client.chat.completions.create(
            model=settings.openai_chat_model,
            messages=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": question}
            ]
        )
        return response.choices[0].message.content


    def search_documents(self, query: str, top_k: int = 5) -> List[RetrievalRef]:
        """Search for relevant document chunks and return with actual content."""
        try:
            # Get query embedding
            query_embedding = self.get_openai_embedding(query)
            
            # Search in Qdrant with lower score threshold
            search_result = self.qdrant_client.search(
                collection_name=settings.qdrant_collection_name,
                query_vector=query_embedding,  # Use default vector field
                limit=top_k,
                score_threshold=0.1  # Lower threshold to get more results
            )
            
            # Convert to RetrievalRef with actual content
            retrieval_refs = []
            for i, hit in enumerate(search_result):
                if hit.payload.get("text", "").strip():
                    retrieval_refs.append(RetrievalRef(
                        doc_id=hit.payload.get("source_file", f"chunk_{i}"),
                        filename=hit.payload.get("source_file", f"chunk_{i}"),
                        page=1,  # Default page
                        chunk_id=hit.id,
                        score=hit.score if hasattr(hit, 'score') else 0.8
                    ))
            
            # If no results with threshold, try without threshold
            if not retrieval_refs:
                search_result = self.qdrant_client.search(
                    collection_name=settings.qdrant_collection_name,
                    query_vector=query_embedding,  # Use default vector field
                    limit=top_k
                )
                
                for i, hit in enumerate(search_result):
                    if hit.payload.get("text", "").strip():
                        retrieval_refs.append(RetrievalRef(
                            doc_id=hit.payload.get("source_file", f"chunk_{i}"),
                            filename=hit.payload.get("source_file", f"chunk_{i}"),
                            page=1,  # Default page
                            chunk_id=hit.id,
                            score=hit.score if hasattr(hit, 'score') else 0.8
                        ))
            
            return retrieval_refs
            
        except Exception as e:
            print(f"Search error: {e}")
            return []

    def get_document_chunks(self, doc_id: str) -> List[dict]:
        """Get all chunks for a specific document."""
        try:
            # Search for documents by source_file in payload
            search_result = self.qdrant_client.scroll(
                collection_name=settings.qdrant_collection_name,
                scroll_filter=Filter(
                    must=[
                        FieldCondition(
                            key="source_file",
                            match=MatchValue(value=doc_id)
                        )
                    ]
                ),
                limit=100
            )
            
            chunks = []
            for point in search_result[0]:
                chunks.append({
                    "chunk_id": point.id,
                    "content": point.payload.get("text", ""),
                    "metadata": point.payload
                })
            
            return chunks
        except Exception as e:
            return []

    def delete_document(self, doc_id: str) -> bool:
        """Delete all chunks for a specific document."""
        try:
            # Validate doc_id
            if not doc_id or doc_id is None:
                return False
            
            # Delete points by source_file filter
            self.qdrant_client.delete(
                collection_name=settings.qdrant_collection_name,
                points_selector=Filter(
                    must=[
                        FieldCondition(
                            key="source_file",
                            match=MatchValue(value=doc_id)
                        )
                    ]
                )
            )
            return True
        except Exception as e:
            return False


# Global RAG engine instance
rag_engine = RAGEngine()
